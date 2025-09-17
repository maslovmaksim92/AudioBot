from fastapi import APIRouter, HTTPException, UploadFile, File, Form, Depends
from pydantic import BaseModel
from typing import List, Dict, Any
import os, io, zipfile, json, logging
from datetime import datetime, timezone, timedelta

from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession, async_sessionmaker
from sqlalchemy.orm import declarative_base
from sqlalchemy import Column, Integer, String, Text, DateTime, ForeignKey
from sqlalchemy import text as sa_text

from emergentintegrations.llm.chat import LlmChat, UserMessage
from openai import AsyncOpenAI
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook
import tiktoken

logger = logging.getLogger("ai_training")

router = APIRouter(prefix="/api/ai-knowledge")

DATABASE_URL = os.environ.get('DATABASE_URL', '').strip()
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY','').strip()
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY','').strip()

ALLOWED_EXT = {'.pdf','.docx','.txt','.xlsx','.zip'}
MAX_FILE_MB = int(os.environ.get('AI_MAX_FILE_MB', '50'))
MAX_TOTAL_MB = int(os.environ.get('AI_MAX_TOTAL_MB', '200'))
DEFAULT_TOP_K = 10

Base = declarative_base()
engine = None
AsyncSessionLocal = None

class AIDocument(Base):
    __tablename__ = 'ai_documents'
    id = Column(String, primary_key=True)
    filename = Column(String, nullable=False)
    mime = Column(String)
    size_bytes = Column(Integer)
    summary = Column(Text)
    pages = Column(Integer)
    created_at = Column(DateTime(timezone=True), default=datetime.now(timezone.utc))

class AIChunk(Base):
    __tablename__ = 'ai_chunks'
    id = Column(String, primary_key=True)
    document_id = Column(String, ForeignKey('ai_documents.id', ondelete='CASCADE'), index=True, nullable=False)
    chunk_index = Column(Integer, nullable=False)
    content = Column(Text, nullable=False)
    # embedding column created by Alembic as Vector(3072)

class AIUploadTemp(Base):
    __tablename__ = 'ai_uploads_temp'
    upload_id = Column(String, primary_key=True)
    meta = Column(Text, nullable=False)  # JSONB in DB; we cast on insert
    expires_at = Column(DateTime(timezone=True), nullable=False)

async def init_db():
    global engine, AsyncSessionLocal
    if not DATABASE_URL:
        logger.warning('DATABASE_URL is not configured; AI DB features disabled')
        return
    if engine is not None:
        return
    engine = create_async_engine(DATABASE_URL, echo=False, pool_pre_ping=True, future=True)
    AsyncSessionLocal = async_sessionmaker(bind=engine, class_=AsyncSession, expire_on_commit=False)
    # Apply Alembic migrations
    try:
        import subprocess, pathlib
        here = pathlib.Path(__file__).parent
        alembic_ini = str((here / 'alembic.ini').resolve())
        if not os.path.exists(alembic_ini):
            alembic_ini = str((here.parent / 'alembic.ini').resolve())
        subprocess.run(['alembic', '-c', alembic_ini, 'upgrade', 'head'], check=False)
        logger.info('Alembic migrations executed for AI')
    except Exception as e:
        logger.warning(f'Alembic run error (AI): {e}')

async def get_db():
    if AsyncSessionLocal is None:
        raise HTTPException(status_code=500, detail='Database is not initialized')
    async with AsyncSessionLocal() as session:
        yield session

async def _split_into_chunks(text: str, target_tokens: int = 1200, overlap: int = 200) -> List[str]:
    enc = tiktoken.get_encoding('cl100k_base')
    toks = enc.encode(text)
    chunks = []
    i = 0
    while i < len(toks):
        window = toks[i:i+target_tokens]
        chunks.append(enc.decode(window))
        i += max(1, target_tokens - overlap)
    return chunks

async def _embed_texts(texts: List[str]) -> List[List[float]]:
    if not OPENAI_API_KEY:
        return [[0.0]*3072 for _ in texts]
    client = AsyncOpenAI(api_key=OPENAI_API_KEY)
    out = []
    for t in texts:
        try:
            r = await client.embeddings.create(model='text-embedding-3-large', input=t)
            out.append(r.data[0].embedding)
        except Exception as e:
            logger.error(f"Embedding error: {e}")
            out.append([0.0]*3072)
    return out

async def _summarize(text: str) -> str:
    if not EMERGENT_LLM_KEY:
        return text[:500]
    try:
        chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=f"ai_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}", system_message="Ты помощник для создания кратких саммари документов.").with_model('openai','gpt-4o-mini')
        prompt = f"Суммируй кратко ключевые пункты (до 120 слов):\n{text[:6000]}"
        resp = await chat.send_message(UserMessage(text=prompt))
        return resp or text[:500]
    except Exception as e:
        logger.warning(f"Summary error: {e}")
        return text[:500]

@router.post('/upload')
async def ai_upload(files: List[UploadFile] = File(...), chunk_tokens: int = Form(1200), overlap: int = Form(200)):
    await init_db()
    if not files:
        raise HTTPException(status_code=400, detail='Нет файлов')
    if AsyncSessionLocal is None:
        raise HTTPException(status_code=500, detail='Database is not initialized')

    all_text = ''
    total_size = 0
    total_pages = 0
    file_stats = []
    for f in files:
        ext = os.path.splitext(f.filename or '')[1].lower()
        data = await f.read()
        size_bytes = len(data)
        if size_bytes > MAX_FILE_MB*1024*1024:
            raise HTTPException(status_code=413, detail=f"Файл {f.filename} превышает {MAX_FILE_MB}MB")
        total_size += size_bytes
        pages = None
        text = ''
        try:
            if ext == '.txt':
                text = data.decode('utf-8', errors='ignore')
            elif ext == '.pdf':
                reader = PdfReader(io.BytesIO(data))
                pages = len(reader.pages)
                parts = []
                for p in reader.pages:
                    try: parts.append(p.extract_text() or '')
                    except Exception: continue
                text = '\n'.join(parts)
            elif ext == '.docx':
                doc = DocxDocument(io.BytesIO(data))
                text = '\n'.join(p.text for p in doc.paragraphs)
            elif ext == '.xlsx':
                wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
                rows = 0; parts = []
                for ws in wb.worksheets:
                    for row in ws.iter_rows(values_only=True):
                        line = ' '.join([str(c) for c in row if c is not None])
                        if line: parts.append(line)
                        rows += 1
                pages = rows
                text = '\n'.join(parts)
            elif ext == '.zip':
                try:
                    zf = zipfile.ZipFile(io.BytesIO(data))
                    for name in zf.namelist():
                        if name.endswith('/'): continue
                        sub_ext = os.path.splitext(name)[1].lower()
                        if sub_ext not in {'.pdf','.docx','.txt','.xlsx'}: continue
                        if os.path.isabs(name) or '..' in name: continue
                        with zf.open(name) as zf_f:
                            sub_raw = zf_f.read()
                            sub_size = len(sub_raw)
                            sub_pages = None
                            sub_text = ''
                            try:
                                if sub_ext == '.pdf':
                                    sreader = PdfReader(io.BytesIO(sub_raw))
                                    sub_pages = len(sreader.pages)
                                    sparts = []
                                    for sp in sreader.pages:
                                        try: sparts.append(sp.extract_text() or '')
                                        except Exception: continue
                                    sub_text = '\n'.join(sparts)
                                elif sub_ext == '.docx':
                                    sdoc = DocxDocument(io.BytesIO(sub_raw))
                                    sub_text = '\n'.join(p.text for p in sdoc.paragraphs)
                                elif sub_ext == '.xlsx':
                                    swb = load_workbook(io.BytesIO(sub_raw), read_only=True, data_only=True)
                                    srows = 0; sparts = []
                                    for ws in swb.worksheets:
                                        for row in ws.iter_rows(values_only=True):
                                            line = ' '.join([str(c) for c in row if c is not None])
                                            if line: sparts.append(line)
                                            srows += 1
                                    sub_pages = srows
                                    sub_text = '\n'.join(sparts)
                                elif sub_ext == '.txt':
                                    sub_text = sub_raw.decode('utf-8', errors='ignore')
                            except Exception: pass
                            file_stats.append({'name': name, 'ext': sub_ext, 'size_bytes': sub_size, 'pages': sub_pages, 'text_chars': len(sub_text or '')})
                            all_text += (sub_text or '') + '\n\n'
                except Exception: pass
                continue
        except Exception:
            text = ''
        file_stats.append({'name': f.filename, 'ext': ext, 'size_bytes': size_bytes, 'pages': pages, 'text_chars': len(text or '')})
        total_pages += (pages or 0)
        all_text += (text or '') + '\n\n'

    preview = ''
    chunks: List[str] = []
    if all_text.strip():
        chunks = await _split_into_chunks(all_text, target_tokens=int(chunk_tokens), overlap=int(overlap))
        s = await _summarize(all_text)
        preview = (s or all_text)[:500]
    else:
        preview = 'Текст не извлечён. Проверьте, что файлы содержат текст (не только изображения).'

    upload_id = os.urandom(8).hex()
    async with AsyncSessionLocal() as s:
        meta = {'filenames':[f.filename for f in files],'summary':preview,'chunks':chunks,'chunks_count':len(chunks),'total_size_bytes':total_size,'total_pages':total_pages,'file_stats':file_stats,'chunk_tokens':int(chunk_tokens),'overlap':int(overlap)}
        await s.execute(sa_text('INSERT INTO ai_uploads_temp (upload_id, meta, expires_at) VALUES (:id, :m::jsonb, :exp)'), {"id": upload_id, "m": json.dumps(meta, ensure_ascii=False), "exp": datetime.now(timezone.utc)+timedelta(hours=6)})
        await s.commit()

    return {'upload_id': upload_id, 'preview': preview, 'chunks': len(chunks), 'stats': {'total_size_bytes': total_size, 'total_pages': total_pages, 'file_stats': file_stats}}

@router.post('/save')
async def ai_save(upload_id: str = Form(...), filename: str = Form('document.txt'), db: AsyncSession = Depends(get_db)):
    row = (await db.execute(sa_text('SELECT meta FROM ai_uploads_temp WHERE upload_id=:id'), {"id": upload_id})).first()
    if not row:
        raise HTTPException(status_code=404, detail='upload_id не найден или истёк')
    raw_meta = row[0]
    meta = json.loads(raw_meta) if isinstance(raw_meta, str) else raw_meta
    chunks = meta.get('chunks', [])
    vectors = await _embed_texts(chunks)
    doc_id = os.urandom(8).hex()
    summary = meta.get('summary') or 'См. превью при загрузке'
    size_bytes = meta.get('total_size_bytes')
    pages = meta.get('total_pages')
    await db.execute(sa_text('INSERT INTO ai_documents (id, filename, mime, size_bytes, summary, pages, created_at) VALUES (:i,:fn,:mime,:sz,:sm,:pg,:ca)'), {"i": doc_id, "fn": filename, "mime": "text/plain", "sz": size_bytes, "sm": (summary[:500] if isinstance(summary,str) else None), "pg": pages, "ca": datetime.now(timezone.utc)})
    for idx, (text, v) in enumerate(zip(chunks, vectors)):
        await db.execute(sa_text('INSERT INTO ai_chunks (id, document_id, chunk_index, content, embedding) VALUES (:i,:d,:x,:c,:e)'), {"i": os.urandom(8).hex(), "d": doc_id, "x": idx, "c": text, "e": v})
    await db.execute(sa_text('DELETE FROM ai_uploads_temp WHERE upload_id=:id'), {"id": upload_id})
    await db.commit()
    return {"document_id": doc_id}

@router.get('/documents')
async def ai_docs_list(db: AsyncSession = Depends(get_db)):
    rows = (await db.execute(sa_text('''
        SELECT d.id, d.filename, d.mime, d.size_bytes, d.summary, d.created_at, d.pages,
               (SELECT COUNT(1) FROM ai_chunks c WHERE c.document_id=d.id) AS chunks_count
        FROM ai_documents d
        ORDER BY d.created_at DESC
        LIMIT 200
    '''))).all()
    docs = []
    for r in rows:
        rid, filename, mime, size_bytes, summary, created_at, pages, chunks_count = r
        docs.append({"id": rid, "filename": filename, "mime": mime, "size_bytes": size_bytes, "summary": summary, "created_at": created_at.isoformat() if created_at else None, "pages": pages, "chunks_count": chunks_count})
    return {"documents": docs}

class SearchRequest(BaseModel):
    query: str
    top_k: int = DEFAULT_TOP_K

@router.post('/search')
async def ai_search(req: SearchRequest, db: AsyncSession = Depends(get_db)):
    if not req.query.strip():
        raise HTTPException(status_code=400, detail='query пуст')
    qv = (await _embed_texts([req.query]))[0]
    rows = (await db.execute(sa_text('''
        SELECT c.document_id, c.chunk_index, c.content, 1 - (c.embedding <=> :qv) as score, d.filename
        FROM ai_chunks c JOIN ai_documents d ON d.id = c.document_id
        ORDER BY c.embedding <=> :qv
        LIMIT :k
    '''), {"qv": qv, "k": req.top_k})).all()
    results = []
    for r in rows:
        doc_id, idx, content, score, filename = r
        results.append({"document_id": doc_id, "chunk_index": idx, "content": content, "score": float(score), "filename": filename})
    return {"results": results}

@router.delete('/document/{doc_id}')
async def ai_delete(doc_id: str, db: AsyncSession = Depends(get_db)):
    await db.execute(sa_text('DELETE FROM ai_chunks WHERE document_id=:id'), {"id": doc_id})
    await db.execute(sa_text('DELETE FROM ai_documents WHERE id=:id'), {"id": doc_id})
    await db.commit()
    return {"ok": True}