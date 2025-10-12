import os
import io
import json
import zipfile
import logging
from datetime import datetime, timezone, timedelta
from typing import Any, Dict, List, Optional
from uuid import uuid4

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Body
from pydantic import BaseModel

# Psycopg3 async
try:
    from psycopg_pool import AsyncConnectionPool
    from psycopg.rows import dict_row
    try:
        from psycopg.types.json import Json as PgJson
    except Exception:  # pragma: no cover
        PgJson = None
    PSYCOPG_AVAILABLE = True
except Exception:  # pragma: no cover
    AsyncConnectionPool = None
    dict_row = None
    PgJson = None
    PSYCOPG_AVAILABLE = False

from openai import AsyncOpenAI
import tiktoken
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook

logger = logging.getLogger("ai_knowledge_router")

# Префикс без '/api' — общий стиль: монтируем router в server.py с prefix="/api"
router = APIRouter(prefix="/ai-knowledge", tags=["AI Knowledge"])

# ========= ENV =========
RAW_DATABASE_URL = (os.environ.get("NEON_DATABASE_URL") or os.environ.get("DATABASE_URL") or "").strip()
OPENAI_API_KEY = (os.environ.get("OPENAI_API_KEY") or "").strip()

ALLOWED_EXT = {".pdf", ".docx", ".txt", ".xlsx", ".zip"}

# ========= DSN helpers =========
from urllib.parse import urlparse, parse_qsl, urlencode, urlunparse


def _scrub_ssl_env():
    for k in ("PGSSLMODE", "PGSSL", "PGSSLCERT", "PGSSLKEY", "PGSSLROOTCERT"):
        os.environ.pop(k, None)


def _normalize_db_url_psycopg(url: str) -> str:
    try:
        if not url:
            return url
        url = url.strip().strip("'\"")
        if url.lower().startswith("psql "):
            url = url[5:].strip().strip("'\"")
        # Scheme -> postgresql://
        if url.startswith("postgresql+asyncpg://"):
            url = url.replace("postgresql+asyncpg://", "postgresql://", 1)
        elif url.startswith("postgres://"):
            url = url.replace("postgres://", "postgresql://", 1)
        elif not url.startswith("postgresql://"):
            url = "postgresql://" + url.split("://", 1)[-1]
        u = urlparse(url)
        q = dict(parse_qsl(u.query, keep_blank_values=True))
        # Map ssl=true → sslmode=require
        if "ssl" in q:
            sval = str(q.get("ssl")).lower()
            q.pop("ssl", None)
            if sval in ("true", "1", "yes", "on"):
                q["sslmode"] = "require"
            elif sval in ("false", "0", "no", "off"):
                q["sslmode"] = "disable"
        allowed = {"disable", "allow", "prefer", "require", "verify-ca", "verify-full"}
        if "sslmode" not in q or str(q["sslmode"]).lower() not in allowed:
            q["sslmode"] = "require"
        # keepalive / timeout / CA bundle
        q.setdefault("keepalives", "1")
        q.setdefault("keepalives_idle", "30")
        q.setdefault("keepalives_interval", "10")
        q.setdefault("keepalives_count", "5")
        q.setdefault("connect_timeout", "10")
        q.setdefault("sslrootcert", "/etc/ssl/certs/ca-certificates.crt")
        new_query = urlencode(q, doseq=True)
        return urlunparse((u.scheme, u.netloc, u.path, u.params, new_query, u.fragment))
    except Exception:
        return url


_scrub_ssl_env()
DATABASE_URL = _normalize_db_url_psycopg(RAW_DATABASE_URL)

# ========= Pool (lazy) =========
pg_pool: Optional[AsyncConnectionPool] = None


async def _ensure_pool() -> bool:
    global pg_pool
    if not PSYCOPG_AVAILABLE or not DATABASE_URL:
        return False
    try:
        if pg_pool is None:
            # Create closed and open explicitly (avoids deprecation warning)
            max_size = int(os.environ.get("AI_PG_MAX_POOL", "5"))
            pg_pool = AsyncConnectionPool(conninfo=DATABASE_URL, min_size=1, max_size=max_size, open=False)
        await pg_pool.open()
        return True
    except Exception as e:
        logger.error(f"AI Knowledge: pool open failed: {e}")
        return False


# ========= Utilities =========

def _tokencut(text: str, max_tokens: int) -> str:
    enc = tiktoken.get_encoding("cl100k_base")
    toks = enc.encode(text or "")
    return enc.decode(toks[:max_tokens])


async def _split_into_chunks(text: str, target_tokens: int = 1200, overlap: int = 200) -> List[str]:
    enc = tiktoken.get_encoding("cl100k_base")
    toks = enc.encode(text or "")
    chunks: List[str] = []
    i = 0
    step = max(1, target_tokens - overlap)
    while i < len(toks):
        window = toks[i : i + target_tokens]
        chunks.append(enc.decode(window))
        i += step
    return chunks


async def _summarize(text: str, max_chars: int = 2000) -> str:
    return (text or "")[:max_chars]


async def _detect_vector_dims() -> int:
    if not await _ensure_pool():
        return 1536
    try:
        async with pg_pool.connection() as conn:  # type: ignore
            conn.row_factory = dict_row
            async with conn.cursor() as cur:
                await cur.execute(
                    """
                    SELECT a.atttypmod AS atttypmod
                    FROM pg_attribute a
                    JOIN pg_class c ON c.oid = a.attrelid
                    WHERE c.relname = 'ai_chunks' AND a.attname = 'embedding'
                    LIMIT 1
                    """
                )
                row = await cur.fetchone()
                if row and isinstance(row.get("atttypmod"), int) and row["atttypmod"] > 4:
                    dims = int(row["atttypmod"]) - 4
                    if 1528 <= dims <= 1536:
                        return 1536
                    return dims
    except Exception as e:
        logger.warning(f"Vector dims detection failed: {e}")
    return 1536


async def _embed_texts_dynamic(texts: List[str]) -> List[List[float]]:
    dims = await _detect_vector_dims()
    model = "text-embedding-3-small" if dims <= 1536 else "text-embedding-3-large"
    if not OPENAI_API_KEY:
        return [[0.0] * dims for _ in texts]
    client = AsyncOpenAI(api_key=OPENAI_API_KEY)
    out: List[List[float]] = []
    for t in texts:
        try:
            r = await client.embeddings.create(model=model, input=t)
            vec = r.data[0].embedding
            if len(vec) != dims:
                if len(vec) > dims:
                    vec = vec[:dims]
                else:
                    vec = vec + [0.0] * (dims - len(vec))
            out.append(vec)
        except Exception as e:
            logger.error(f"Embedding error: {e}")
            out.append([0.0] * dims)
    return out


# ========= Diagnostics =========
class DbCheckResponse(BaseModel):
    connected: bool
    pgvector_available: bool
    pgvector_installed: bool
    ai_tables: List[str]
    embedding_dims: Optional[int] = None
    errors: List[str] = []


@router.get("/db-check", response_model=DbCheckResponse)
async def db_check():
    errors: List[str] = []
    ai_tables: List[str] = []
    connected = False
    pgvector_available = False
    pgvector_installed = False
    dims: Optional[int] = None
    ok = await _ensure_pool()
    if not ok:
        return DbCheckResponse(
            connected=False,
            pgvector_available=False,
            pgvector_installed=False,
            ai_tables=[],
            embedding_dims=None,
            errors=["pool_not_initialized"],
        )
    try:
        async with pg_pool.connection() as conn:  # type: ignore
            conn.row_factory = dict_row
            async with conn.cursor() as cur:
                try:
                    await cur.execute("SELECT 1")
                    await cur.fetchone()
                    connected = True
                except Exception as e:
                    errors.append(f"connect: {e}")
                try:
                    await cur.execute("SELECT default_version FROM pg_available_extensions WHERE name='vector'")
                    row = await cur.fetchone()
                    pgvector_available = bool(row)
                except Exception as e:
                    errors.append(f"available: {e}")
                try:
                    await cur.execute("SELECT extversion FROM pg_extension WHERE extname='vector'")
                    row2 = await cur.fetchone()
                    pgvector_installed = bool(row2)
                except Exception as e:
                    errors.append(f"installed: {e}")
                try:
                    await cur.execute(
                        "SELECT table_name FROM information_schema.tables WHERE table_schema='public' AND table_name LIKE 'ai_%'"
                    )
                    rows = await cur.fetchall()
                    ai_tables = [r["table_name"] for r in rows] if rows else []
                except Exception as e:
                    errors.append(f"ai_tables: {e}")
                try:
                    if "ai_chunks" in ai_tables:
                        await cur.execute(
                            """
                            SELECT a.atttypmod AS atttypmod
                            FROM pg_attribute a JOIN pg_class c ON c.oid = a.attrelid
                            WHERE c.relname = 'ai_chunks' AND a.attname = 'embedding' LIMIT 1
                            """
                        )
                        rowd = await cur.fetchone()
                        if rowd and rowd.get("atttypmod") and rowd["atttypmod"] > 4:
                            _d = int(rowd["atttypmod"]) - 4
                            dims = 1536 if 1528 <= _d <= 1536 else _d
                except Exception as e:
                    errors.append(f"dims: {e}")
    except Exception as e:
        errors.append(f"session: {e}")
    return DbCheckResponse(
        connected=connected,
        pgvector_available=pgvector_available,
        pgvector_installed=pgvector_installed,
        ai_tables=ai_tables,
        embedding_dims=dims,
        errors=errors,
    )


@router.get("/db-dsn")
async def db_dsn():
    raw = RAW_DATABASE_URL
    norm = _normalize_db_url_psycopg(raw)
    def _parse(u: str):
        try:
            p = urlparse(u)
            q = dict(parse_qsl(p.query, keep_blank_values=True))
            host = p.hostname or ""
            port = p.port
            dbname = (p.path or "").lstrip("/")
            username = p.username or ""
            if username:
                username = (username[:2] + "***") if len(username) > 3 else "***"
            return {"scheme": p.scheme, "host": host, "port": port, "database": dbname, "username": username, "query": q}
        except Exception:
            return {"error": "parse_failed"}
    n = _parse(norm) if norm else None
    nq = n and n.get("query") or {}
    normalized = {
        "scheme": n and n.get("scheme"),
        "host": n and n.get("host"),
        "port": n and n.get("port"),
        "database": n and n.get("database"),
        "username": n and n.get("username"),
        "query": f"sslmode={nq.get('sslmode')}" if nq else None,
    }
    return {"raw_present": bool(raw), "normalized": normalized}


# ========= Preview / Study / Status =========
@router.post("/preview")
async def preview(
    file: UploadFile = File(None),
    files: List[UploadFile] = File(None),
    chunk_tokens: int = Form(1200),
    overlap: int = Form(200),
):
    ok = await _ensure_pool()
    if not ok:
        raise HTTPException(status_code=500, detail="Database is not initialized")

    # Support both 'file' and 'files'
    picked: Optional[UploadFile] = None
    if file is not None:
        picked = file
    elif files:
        picked = files[0]
    if not picked:
        raise HTTPException(status_code=400, detail="Файл не передан")

    ext = os.path.splitext(picked.filename or "")[1].lower()
    if ext not in ALLOWED_EXT:
        raise HTTPException(status_code=400, detail=f"Недопустимый формат: {ext}")

    raw = await picked.read()
    text = ""
    pages = None
    try:
        if ext == ".txt":
            text = raw.decode("utf-8", errors="ignore")
        elif ext == ".pdf":
            reader = PdfReader(io.BytesIO(raw))
            pages = len(reader.pages)
            parts = []
            for p in reader.pages:
                try:
                    parts.append(p.extract_text() or "")
                except Exception:
                    continue
            text = "\n".join(parts)
        elif ext == ".docx":
            doc = DocxDocument(io.BytesIO(raw))
            text = "\n".join(p.text for p in doc.paragraphs)
        elif ext == ".xlsx":
            wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
            rows = 0
            parts = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    line = " ".join([str(c) for c in row if c is not None])
                    if line:
                        parts.append(line)
                    rows += 1
            pages = rows
            text = "\n".join(parts)
        elif ext == ".zip":
            out = []
            try:
                zf = zipfile.ZipFile(io.BytesIO(raw))
                for name in zf.namelist():
                    if name.endswith("/"):
                        continue
                    se = os.path.splitext(name)[1].lower()
                    if se not in {".pdf", ".docx", ".txt", ".xlsx"}:
                        continue
                    if os.path.isabs(name) or ".." in name:
                        continue
                    with zf.open(name) as fz:
                        sub = fz.read()
                        try:
                            if se == ".txt":
                                out.append(sub.decode("utf-8", errors="ignore"))
                            elif se == ".pdf":
                                rr = PdfReader(io.BytesIO(sub))
                                pp = []
                                for sp in rr.pages:
                                    try:
                                        pp.append(sp.extract_text() or "")
                                    except Exception:
                                        continue
                                out.append("\n".join(pp))
                            elif se == ".docx":
                                dd = DocxDocument(io.BytesIO(sub))
                                out.append("\n".join(p.text for p in dd.paragraphs))
                            elif se == ".xlsx":
                                swb = load_workbook(io.BytesIO(sub), read_only=True, data_only=True)
                                pp = []
                                for ws in swb.worksheets:
                                    for row in ws.iter_rows(values_only=True):
                                        line = " ".join([str(c) for c in row if c is not None])
                                        if line:
                                            pp.append(line)
                                out.append("\n".join(pp))
                        except Exception:
                            continue
            except Exception:
                pass
            text = "\n\n".join(out)
    except Exception:
        text = ""

    chunks = await _split_into_chunks(text or "", target_tokens=int(chunk_tokens), overlap=int(overlap))
    desc = await _summarize(text or "", max_chars=2000)

    upload_id = str(uuid4())
    meta = {
        "filename": picked.filename,
        "summary": desc,
        "chunks": chunks,
        "chunks_count": len(chunks),
        "size_bytes": len(raw),
        "pages": pages,
        "chunk_tokens": int(chunk_tokens),
        "overlap": int(overlap),
        "status": "ready",
    }
    try:
        async with pg_pool.connection() as conn:  # type: ignore
            async with conn.cursor() as cur:
                payload = PgJson(meta) if PgJson else json.dumps(meta, ensure_ascii=False)
                await cur.execute(
                    "INSERT INTO ai_uploads_temp (upload_id, meta, expires_at) VALUES (%(id)s, %(m)s::jsonb, %(exp)s)",
                    {"id": upload_id, "m": payload, "exp": datetime.now(timezone.utc) + timedelta(hours=6)},
                )
                await conn.commit()
    except Exception as e:
        logger.error(f"DB insert preview error: {e}")
        raise HTTPException(status_code=500, detail="Database write error")

    return {
        "upload_id": upload_id,
        "preview": desc,
        "chunks": len(chunks),
        "stats": {
            "total_size_bytes": len(raw),
            "total_pages": pages or 0,
            "file_stats": [
                {
                    "name": picked.filename,
                    "ext": ext,
                    "size_bytes": len(raw),
                    "pages": pages,
                    "text_chars": len(text or ""),
                }
            ],
        },
    }


class StatusResponse(BaseModel):
    status: str
    detail: Optional[str] = None


@router.get("/status", response_model=StatusResponse)
async def status(upload_id: str):
    ok = await _ensure_pool()
    if not ok:
        raise HTTPException(status_code=500, detail="Database is not initialized")
    try:
        async with pg_pool.connection() as conn:  # type: ignore
            conn.row_factory = dict_row
            async with conn.cursor() as cur:
                await cur.execute("SELECT meta FROM ai_uploads_temp WHERE upload_id=%(id)s", {"id": upload_id})
                row = await cur.fetchone()
                if row:
                    raw_meta = row.get("meta")
                    meta = raw_meta if isinstance(raw_meta, dict) else (json.loads(raw_meta) if isinstance(raw_meta, str) else {})
                    return StatusResponse(status=meta.get("status") or "ready")
    except Exception as e:
        logger.error(f"status error: {e}")
        raise HTTPException(status_code=500, detail="Database read error")
    return StatusResponse(status="done")


@router.post("/study")
async def study(
    upload_id: str = Form(None),
    filename: str = Form(None),
    category: str = Form(None),
    json_body: Optional[Dict[str, Any]] = Body(None),
):
    if (not upload_id or not filename or not category) and json_body:
        upload_id = upload_id or json_body.get("upload_id")
        filename = filename or json_body.get("filename") or "document.txt"
        category = category or json_body.get("category") or "Клининг"
    if not upload_id:
        raise HTTPException(status_code=422, detail="upload_id is required")
    if not filename:
        filename = "document.txt"
    if not category:
        category = "Клининг"
    ok = await _ensure_pool()
    if not ok:
        raise HTTPException(status_code=500, detail="Database is not initialized")

    try:
        async with pg_pool.connection() as conn:  # type: ignore
            conn.row_factory = dict_row
            async with conn.cursor() as cur:
                await cur.execute("SELECT meta FROM ai_uploads_temp WHERE upload_id=%(id)s", {"id": upload_id})
                row = await cur.fetchone()
                if not row:
                    raise HTTPException(status_code=404, detail="upload_id не найден или истёк")
                raw_meta = row.get("meta")
                meta = raw_meta if isinstance(raw_meta, dict) else (json.loads(raw_meta) if isinstance(raw_meta, str) else {})
                chunks: List[str] = meta.get("chunks") or []
                vectors = await _embed_texts_dynamic(chunks)
                # Normalize vector length to 1536 to avoid rare atttypmod anomalies
                norm = []
                for v in vectors:
                    if not isinstance(v, list):
                        v = []
                    if 1528 <= len(v) <= 1536:
                        if len(v) < 1536:
                            v = v + [0.0] * (1536 - len(v))
                        elif len(v) > 1536:
                            v = v[:1536]
                    norm.append(v)
                doc_id = str(uuid4())
                summary = meta.get("summary") or ""
                size_bytes = int(meta.get("size_bytes") or meta.get("total_size_bytes") or 0)
                pages = meta.get("pages") or meta.get("total_pages")
                mime = f"text/plain; category={category}"
                await cur.execute(
                    "INSERT INTO ai_documents (id, filename, mime, size_bytes, summary, pages, created_at) VALUES (%(i)s,%(fn)s,%(mime)s,%(sz)s,%(sm)s,%(pg)s,%(ca)s)",
                    {
                        "i": doc_id,
                        "fn": filename,
                        "mime": mime,
                        "sz": size_bytes,
                        "sm": (summary[:500] if isinstance(summary, str) else None),
                        "pg": pages,
                        "ca": datetime.now(timezone.utc),
                    },
                )
                for idx, (text, v) in enumerate(zip(chunks, norm)):
                    v_str = "[" + ",".join(str(float(x)) for x in (v or [])) + "]"
                    await cur.execute(
                        "INSERT INTO ai_chunks (id, document_id, chunk_index, content, embedding) VALUES (%(i)s,%(d)s,%(x)s,%(c)s,(%(e)s)::vector)",
                        {"i": str(uuid4()), "d": doc_id, "x": idx, "c": text, "e": v_str},
                    )
                await cur.execute("DELETE FROM ai_uploads_temp WHERE upload_id=%(id)s", {"id": upload_id})
                await conn.commit()
        return {"document_id": doc_id, "chunks": len(chunks), "category": category}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"study error: {e}")
        raise HTTPException(status_code=500, detail="Database write error")


# ========= Remember =========
class RememberRequest(BaseModel):
    text: str
    category: Optional[str] = 'Notes'
    filename: Optional[str] = None

@router.post('/remember')
async def remember(req: RememberRequest):
    if not (req and (req.text or '').strip()):
        raise HTTPException(status_code=400, detail='text is required')
    ok = await _ensure_pool()
    if not ok:
        raise HTTPException(status_code=500, detail='Database is not initialized')
    text = (req.text or '').strip()
    try:
        chunks = await _split_into_chunks(text, target_tokens=900, overlap=200)
        vectors = await _embed_texts_dynamic(chunks)
        # Normalize to 1536 if needed
        norm = []
        for v in vectors:
            if not isinstance(v, list):
                v = []
            if 1528 <= len(v) <= 1536:
                if len(v) < 1536:
                    v = v + [0.0] * (1536 - len(v))
                elif len(v) > 1536:
                    v = v[:1536]
            norm.append(v)
        doc_id = str(uuid4())
        filename = req.filename or f'note_{datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")}.txt'
        mime = f'text/plain; category={req.category or "Notes"}'
        async with pg_pool.connection() as conn:  # type: ignore
            async with conn.cursor() as cur:
                await cur.execute(
                    'INSERT INTO ai_documents (id, filename, mime, size_bytes, summary, pages, created_at) VALUES (%(i)s,%(fn)s,%(mime)s,%(sz)s,%(sm)s,%(pg)s,%(ca)s)',
                    { 'i': doc_id, 'fn': filename, 'mime': mime, 'sz': len(text.encode("utf-8")), 'sm': (text[:500] if isinstance(text,str) else None), 'pg': None, 'ca': datetime.now(timezone.utc) }
                )
                for idx, (c, v) in enumerate(zip(chunks, norm)):
                    v_str = '[' + ','.join(str(float(x)) for x in (v or [])) + ']'
                    await cur.execute(
                        'INSERT INTO ai_chunks (id, document_id, chunk_index, content, embedding) VALUES (%(i)s,%(d)s,%(x)s,%(c)s,(%(e)s)::vector)',
                        { 'i': str(uuid4()), 'd': doc_id, 'x': idx, 'c': c, 'e': v_str }
                    )
                await conn.commit()
        return { 'ok': True, 'document_id': doc_id, 'chunks': len(chunks) }
    except Exception as e:
        logger.error(f'remember error: {e}')
        raise HTTPException(status_code=500, detail='Database write error')

# ========= Search / Answer =========
class SearchRequest(BaseModel):
    query: str
    top_k: int = 10


@router.post("/search")
async def search(req: SearchRequest):
    if not await _ensure_pool():
        return {"results": []}
    q = (req.query or "").strip()
    if not q:
        return {"results": []}
    try:
        qvec = (await _embed_texts_dynamic([q]))[0]
        qvec_str = "[" + ",".join(str(float(x)) for x in qvec) + "]"
        async with pg_pool.connection() as conn:  # type: ignore
            conn.row_factory = dict_row
            async with conn.cursor() as cur:
                await cur.execute(
                    """
                    SELECT d.id as document_id, d.filename, c.chunk_index, c.content,
                           1 - (c.embedding &lt;=&gt; (%(qv)s)::vector) as score
                    FROM ai_chunks c
                    JOIN ai_documents d ON d.id = c.document_id
                    ORDER BY c.embedding &lt;=&gt; (%(qv)s)::vector
                    LIMIT %(k)s
                    """,
                    {"qv": qvec_str, "k": int(req.top_k)},
                )
                rows = await cur.fetchall()
                results = []
                for r in rows or []:
                    results.append(
                        {
                            "document_id": r.get("document_id"),
                            "chunk_index": r.get("chunk_index"),
                            "content": r.get("content"),
                            "score": float(r.get("score") or 0.0),
                            "filename": r.get("filename"),
                        }
                    )
                return {"results": results}
    except Exception as e:
        logger.error(f"search error: {e}")
        return {"results": []}


class AnswerRequest(BaseModel):
    question: str
    top_k: int = 8
    category: Optional[str] = None
    min_score: float = 0.05
    session_id: Optional[str] = None


@router.post("/answer")
async def answer(req: AnswerRequest):
    q = (req.question or "").strip()
    if not q:
        return {"answer": "", "citations": []}
    citations: List[Dict[str, Any]] = []
    context_blocks: List[str] = []
    try:
        # 1) Пытаемся найти контекст в БЗ
        if await _ensure_pool():
            qvec = (await _embed_texts_dynamic([q]))[0]
            qvec_str = "[" + ",".join(str(float(x)) for x in qvec) + "]"
            async with pg_pool.connection() as conn:  # type: ignore
                conn.row_factory = dict_row
                async with conn.cursor() as cur:
                    cat_clause = ""
                    params = {"qv": qvec_str, "k": int(req.top_k)}
                    if req.category:
                        cat_clause = "WHERE d.mime ILIKE %(cat)s"
                        params["cat"] = f"%category={req.category}%"
                    await cur.execute(
                        f"""
                        SELECT d.id as document_id, d.filename, c.chunk_index, c.content,
                               1 - (c.embedding &lt;=&gt; (%(qv)s)::vector) as score
                        FROM ai_chunks c
                        JOIN ai_documents d ON d.id = c.document_id
                        {cat_clause}
                        ORDER BY c.embedding &lt;=&gt; (%(qv)s)::vector
                        LIMIT %(k)s
                        """,
                        params,
                    )
                    rows = await cur.fetchall()
                    prepared: List[Dict[str, Any]] = []
                    for r in rows or []:
                        prepared.append(
                            {
                                "document_id": r.get("document_id"),
                                "filename": r.get("filename"),
                                "chunk_index": r.get("chunk_index"),
                                "score": float(r.get("score") or 0.0),
                                "content": r.get("content") or "",
                            }
                        )
                    prepared.sort(key=lambda a: a["score"], reverse=True)
                    used = 0
                    for r in prepared:
                        if r["score"] < float(req.min_score):
                            continue
                        excerpt = r["content"][:600]
                        citations.append({
                            "document_id": r["document_id"],
                            "filename": r["filename"],
                            "chunk_index": r["chunk_index"],
                            "score": r["score"],
                            "excerpt": excerpt,
                        })
                        context_blocks.append(f"Источник: {r['filename']} (фрагм. #{r['chunk_index']})\n{excerpt}")
                        used += 1
                        if used >= int(req.top_k):
                            break
                    # Если все ниже порога — берём топ-1 при минимальном score > 0.02
                    if used == 0 and prepared:
                        top = prepared[0]
                        if top["score"] > 0.02:
                            excerpt = top["content"][:600]
                            citations.append({
                                "document_id": top["document_id"],
                                "filename": top["filename"],
                                "chunk_index": top["chunk_index"],
                                "score": top["score"],
                                "excerpt": excerpt,
                            })
                            context_blocks.append(f"Источник: {top['filename']} (фрагм. #{top['chunk_index']})\n{excerpt}")
                            used = 1
        context_text = ("\n\n".join(context_blocks))[:6000]

        # 2) Генерация ответа напрямую через OpenAI (если ключ доступен)
        if OPENAI_API_KEY:
            client = AsyncOpenAI(api_key=OPENAI_API_KEY)
            system_prompt = (
                "Ты ассистент VasDom. Отвечай на русском, кратко и по делу. "
                "Если есть контекст из базы знаний — опирайся ТОЛЬКО на него, без выдумок."
            )
            user_prompt = (
                f"Вопрос: {q}\n\n" + (f"Контекст (используй только его):\n{context_text}\n\n" if context_text else "") + "Ответ:"
            )
            try:
                resp = await client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    temperature=0.3,
                    max_tokens=400,
                )
                answer_text = (resp.choices[0].message.content or "").strip()
                return {"answer": answer_text, "citations": citations}
            except Exception as e:
                logger.error(f"OpenAI chat error: {e}")
                # мягкий фолбэк
                if context_text:
                    fallback = "\n\n".join(context_blocks[:2])
                    return {"answer": fallback, "citations": citations}
                return {"answer": "Готов помочь по рабочим вопросам.", "citations": []}
        else:
            # Без LLM: возвращаем короткий fallback
            if context_text:
                fallback = "\n\n".join(context_blocks[:2])
                return {"answer": fallback, "citations": citations}
            return {"answer": "Готов помочь по рабочим вопросам.", "citations": []}
    except Exception as e:
        logger.error(f"answer error: {e}")
        # мягкий фолбэк
        return {"answer": "Готов помочь по рабочим вопросам.", "citations": []}


# ========= Feedback =========
class FeedbackRequest(BaseModel):
    channel: str  # 'web' | 'telegram'
    rating: int
    question: Optional[str] = None
    answer: Optional[str] = None
    comment: Optional[str] = None
    message_id: Optional[str] = None
    chat_id: Optional[str] = None
    user_id: Optional[str] = None


@router.post("/feedback")
async def feedback(req: FeedbackRequest):
    ok = await _ensure_pool()
    if not ok:
        # Логируем хотя бы в файл
        logger.info(f"feedback (no-db): {req.model_dump()}")
        return {"ok": True}
    try:
        async with pg_pool.connection() as conn:  # type: ignore
            async with conn.cursor() as cur:
                await cur.execute(
                    """
                    CREATE TABLE IF NOT EXISTS ai_feedback (
                        id uuid primary key,
                        channel text,
                        message_id text,
                        chat_id text,
                        user_id text,
                        rating int,
                        comment text,
                        question text,
                        answer text,
                        created_at timestamptz default now()
                    )
                    """
                )
                await cur.execute(
                    """
                    INSERT INTO ai_feedback (id, channel, message_id, chat_id, user_id, rating, comment, question, answer)
                    VALUES (%(i)s, %(ch)s, %(mid)s, %(cid)s, %(uid)s, %(r)s, %(c)s, %(q)s, %(a)s)
                    """,
                    {
                        "i": str(uuid4()),
                        "ch": req.channel,
                        "mid": req.message_id,
                        "cid": req.chat_id,
                        "uid": req.user_id,
                        "r": int(req.rating),
                        "c": (req.comment or "")[:2000],
                        "q": (req.question or "")[:4000],
                        "a": (req.answer or "")[:4000],
                    },
                )
                await conn.commit()
        return {"ok": True}
    except Exception as e:
        logger.error(f"feedback error: {e}")
        return {"ok": False}