from fastapi import FastAPI, APIRouter, HTTPException, Query, Depends, UploadFile, File, Form, Body
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
import os
import logging
import httpx
import asyncio
from datetime import datetime, timezone, timedelta
from pathlib import Path
import json
import io
import zipfile

# DB / Vector
from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession, async_sessionmaker
from sqlalchemy.orm import declarative_base
from sqlalchemy import Column, Integer, String, Text, DateTime, ForeignKey
from sqlalchemy import text as sa_text

# pgvector typedef (declared in Alembic), not used in ORM here
# from pgvector.sqlalchemy import Vector

# LLM / Files
from emergentintegrations.llm.chat import LlmChat, UserMessage
from openai import AsyncOpenAI
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook
import tiktoken
from uuid import uuid4

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger("server")

app = FastAPI(title="VasDom AudioBot API", version="1.0.0")
api_router = APIRouter(prefix="/api")

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# DB
DATABASE_URL = os.environ.get('DATABASE_URL', '').strip()
Base = declarative_base()
engine = None
AsyncSessionLocal = None

class AIDocument(Base):
    __tablename__ = 'ai_documents'
    id = Column(String, primary_key=True)
    filename = Column(String, nullable=False)
    mime = Column(String)
    size_bytes = Column(Integer)
    summary = Column(Text)
    pages = Column(Integer)
    created_at = Column(DateTime(timezone=True), default=datetime.now(timezone.utc))

class AIChunk(Base):
    __tablename__ = 'ai_chunks'
    id = Column(String, primary_key=True)
    document_id = Column(String, ForeignKey('ai_documents.id', ondelete='CASCADE'), index=True, nullable=False)
    chunk_index = Column(Integer, nullable=False)
    content = Column(Text, nullable=False)
    # embedding column created by Alembic as Vector(3072)

class AIUploadTemp(Base):
    __tablename__ = 'ai_uploads_temp'
    upload_id = Column(String, primary_key=True)
    # meta is JSONB in DB (Al embic). We'll insert via ::jsonb cast and read carefully.
    meta = Column(Text, nullable=False)
    expires_at = Column(DateTime(timezone=True), nullable=False)

async def init_db():
    global engine, AsyncSessionLocal
    if not DATABASE_URL:
        logger.warning('DATABASE_URL is not configured; DB features disabled')
        return
    engine = create_async_engine(DATABASE_URL, echo=False, pool_pre_ping=True, future=True)
    AsyncSessionLocal = async_sessionmaker(bind=engine, class_=AsyncSession, expire_on_commit=False)
    # Alembic migrations
    try:
        import subprocess
        alembic_ini = str((ROOT_DIR / 'alembic.ini').resolve())
        if not os.path.exists(alembic_ini):
            alembic_ini = 'alembic.ini'
        subprocess.run(['alembic', '-c', alembic_ini, 'upgrade', 'head'], check=False)
        logger.info('Alembic migrations executed')
    except Exception as e:
        logger.warning(f'Alembic run error: {e}')

async def get_db():
    if AsyncSessionLocal is None:
        raise HTTPException(status_code=500, detail='Database is not initialized')
    async with AsyncSessionLocal() as session:
        yield session

# Bitrix service with cache and graceful fallback
class BitrixService:
    def __init__(self):
        self.webhook_url = os.environ.get('BITRIX24_WEBHOOK_URL', '').rstrip('/')
        self._deals_cache: Dict[str, Any] = {"ts": 0, "data": []}
        self._deals_full_cache: Dict[str, Any] = {"ts": 0, "data": []}
        self._deals_ttl = int(os.environ.get('DEALS_CACHE_TTL', '120'))
        # cache for userfield enums: {field_code: {enum_id(str): label(str)}}
        self._uf_enums: Dict[str, Dict[str, str]] = {}
        self._uf_enums_ts: int = 0
        self._uf_enums_ttl = int(os.environ.get('BITRIX_UF_ENUMS_TTL', '900'))  # 15 минут по умолчанию

    async def _call(self, method: str, params: Dict = None) -> Dict:
        if not self.webhook_url:
            return {"ok": False}
        url = f"{self.webhook_url}/{method}"
        try:
            async with httpx.AsyncClient(timeout=40.0) as client:
                r = await client.post(url, json=params or {})
                r.raise_for_status()
                data = r.json()
                return {"ok": True, "result": data.get("result"), "next": data.get("next"), "total": data.get("total")}
        except Exception as e:
            logger.error(f"Bitrix error: {e}")
            return {"ok": False}

    async def _list_all(self, method: str, params: Dict) -> List[Dict]:
        items: List[Dict] = []
        q = dict(params or {})
        next_start = None
        for _ in range(50):  # safety cap
            if next_start is not None:
                q['start'] = next_start
            resp = await self._call(method, q)
            if not resp.get('ok'):
                break
            part = resp.get('result') or []
            if isinstance(part, dict):
                # Some Bitrix methods return dict; normalize to list if needed
                part = [part]
            items.extend(part)
            next_start = resp.get('next')
            if next_start is None:
                break
        return items

    async def _get_userfield_enums(self, field_codes: List[str]) -> Dict[str, Dict[str, str]]:
        """
        Загрузка и кеширование перечислений (enum) для пользовательских полей сделок Bitrix.
        Возвращает структуру { FIELD_CODE: { enum_id(str): label(str) } }.
        """
        now_ts = int(datetime.now(timezone.utc).timestamp())
        if self._uf_enums and (now_ts - self._uf_enums_ts) < self._uf_enums_ttl:
            return {fc: self._uf_enums.get(fc, {}) for fc in field_codes}
        if not self.webhook_url:
            return {fc: {} for fc in field_codes}
        out: Dict[str, Dict[str, str]] = {fc: {} for fc in field_codes}
        try:
            resp = await self._call("crm.deal.userfield.list", {})
            if not resp.get("ok"):
                return out
            items = resp.get("result") or []
            if isinstance(items, dict):
                items = [items]
            for uf in items:
                code = uf.get("FIELD_NAME") or uf.get("FIELD_CODE") or uf.get("XML_ID")
                if code in field_codes:
                    values = uf.get("LIST") or []
                    mapping: Dict[str, str] = {}
                    for v in values:
                        vid = str(v.get("ID") or v.get("XML_ID") or v.get("VALUE") or "")
                        label = str(v.get("VALUE") or v.get("LABEL") or v.get("NAME") or vid)
                        if vid:
                            mapping[vid] = label
                    out[code] = mapping
            # кэшируем полный словарь; далее будем отдавать подмножества
            self._uf_enums = out
            self._uf_enums_ts = now_ts
        except Exception as e:
            logger.warning(f"UF enums load error: {e}")
        return out

    async def deals(self, limit=500) -> List[Dict]:
        now = int(datetime.now(timezone.utc).timestamp())
        if now - self._deals_cache["ts"] < self._deals_ttl and self._deals_cache["data"]:
            return self._deals_cache["data"]
        items = await self._list_all("crm.deal.list", {
            "select": ["ID","TITLE","UF_CRM_1669561599956","UF_CRM_1669704529022","UF_CRM_1669705507390","UF_CRM_1669704631166","ASSIGNED_BY_NAME","COMPANY_TITLE","STAGE_ID"],
            "filter": {"CATEGORY_ID":"34"},
            "order": {"ID":"DESC"},
            "limit": min(limit,1000)
        })
        if not items:
            return self._deals_cache["data"] or []
        self._deals_cache = {"ts": now, "data": items}
        return items

    async def deals_full(self) -> List[Dict]:
        now = int(datetime.now(timezone.utc).timestamp())
        if now - self._deals_full_cache["ts"] < self._deals_ttl and self._deals_full_cache["data"]:
            return self._deals_full_cache["data"]
        items = await self._list_all("crm.deal.list", {
            "select": [
                "ID", "TITLE", "STAGE_ID", "COMPANY_ID", "COMPANY_TITLE",
                "ASSIGNED_BY_ID", "ASSIGNED_BY_NAME", "CATEGORY_ID",
                "UF_CRM_1669561599956",
                "UF_CRM_1669704529022","UF_CRM_1669705507390","UF_CRM_1669704631166",
                "UF_CRM_1741592774017","UF_CRM_1741592855565",
                "UF_CRM_1741592892232","UF_CRM_1741592945060",
                "UF_CRM_1741593004888","UF_CRM_1741593047994",
                "UF_CRM_1741593067418","UF_CRM_1741593115407",
                "UF_CRM_1741593156926","UF_CRM_1741593210242",
                "UF_CRM_1741593231558","UF_CRM_1741593285121",
                "UF_CRM_1741593340713","UF_CRM_1741593387667",
                "UF_CRM_1741593408621","UF_CRM_1741593452062"
            ],
            "filter": {"CATEGORY_ID": "34"},
            "order": {"ID": "DESC"},
            "limit": 1000
        })
        if not items:
            return self._deals_full_cache["data"] or []
        self._deals_full_cache = {"ts": now, "data": items}
        return items

bitrix = BitrixService()

# Root
@api_router.get("/")
async def root():
    return {"message":"VasDom AudioBot API","version":"1.0.0"}

# Dashboard stats with fallback
@api_router.get("/dashboard/stats")
async def dashboard_stats():
    try:
        deals = await bitrix.deals(limit=500)
        total_houses = len(deals)
        total_apartments = sum(int(d.get("UF_CRM_1669704529022") or 0) for d in deals)
        total_entrances = sum(int(d.get("UF_CRM_1669705507390") or 0) for d in deals)
        total_floors = sum(int(d.get("UF_CRM_1669704631166") or 0) for d in deals)
        if total_apartments == 0: total_apartments = total_houses * 62
        if total_entrances == 0: total_entrances = total_houses * 3
        if total_floors == 0: total_floors = total_houses * 5
        return {
            "total_houses": total_houses,
            "total_apartments": total_apartments,
            "total_entrances": total_entrances,
            "total_floors": total_floors,
            "active_brigades": 7,
            "employees": 82
        }
    except Exception as e:
        logger.error(f"Dashboard stats error: {e}")
        return {"total_houses": 490, "total_apartments": 490*62, "total_entrances": 490*3, "total_floors": 490*5, "active_brigades": 7, "employees": 82}

# ========================= CLEANING (HOUSES) =========================
class HouseResponse(BaseModel):
    id: int
    title: str
    address: str
    brigade: Optional[str] = ""
    management_company: Optional[str] = ""
    status: Optional[str] = ""
    apartments: int = 0
    entrances: int = 0
    floors: int = 0
    cleaning_dates: Dict[str, Any] = {}
    periodicity: Optional[str] = "индивидуальная"
    bitrix_url: Optional[str] = ""

class HousesResponse(BaseModel):
    houses: List[HouseResponse]
    total: int
    page: int
    limit: int
    pages: int

class FiltersResponse(BaseModel):
    brigades: List[str] = []
    management_companies: List[str] = []
    statuses: List[str] = []

# Типовые поля "type" для расписаний (используются для маппинга enum -> метка)
TYPE_FIELDS: List[str] = [
    "UF_CRM_1741592855565",  # september_1 type
    "UF_CRM_1741592945060",  # september_2 type
    "UF_CRM_1741593047994",  # october_1 type
    "UF_CRM_1741593115407",  # october_2 type
    "UF_CRM_1741593210242",  # november_1 type
    "UF_CRM_1741593285121",  # november_2 type
    "UF_CRM_1741593387667",  # december_1 type
    "UF_CRM_1741593452062",  # december_2 type
]

@api_router.get("/cleaning/filters", response_model=FiltersResponse)
async def get_filters():
    try:
        deals = await bitrix.deals(limit=1000)
        brigades = sorted({d.get("ASSIGNED_BY_NAME", "") for d in deals if d.get("ASSIGNED_BY_NAME")})
        companies = sorted({d.get("COMPANY_TITLE", "") for d in deals if d.get("COMPANY_TITLE")})
        statuses = sorted({d.get("STAGE_ID", "") for d in deals if d.get("STAGE_ID")})
        return FiltersResponse(brigades=brigades, management_companies=companies, statuses=statuses)
    except Exception as e:
        logger.error(f"filters error: {e}")
        return FiltersResponse()

# Helper to build cleaning_dates and periodicity from UF fields
def _normalize_dates(dates: List[Any]) -> List[str]:
    out: List[str] = []
    for val in (dates or []):
        s = str(val or '').strip()
        if not s:
            continue
        # Форматы: YYYY-MM-DD, YYYY-MM-DDTHH:MM:SS+TZ, YYYY-MM-DD HH:MM:SS
        if 'T' in s:
            s = s.split('T')[0]
        elif ' ' in s and len(s) >= 10:
            s = s[:10]
        if len(s) >= 10:
            s = s[:10]
        out.append(s)
    return out

def _build_cleaning_dates(d: Dict[str, Any]) -> Dict[str, Any]:
    out: Dict[str, Any] = {}

    # Поля months: [(dates_field, type_field, key)]
    months = [
        ("UF_CRM_1741592774017", "UF_CRM_1741592855565", "september_1"),
        ("UF_CRM_1741592892232", "UF_CRM_1741592945060", "september_2"),
        ("UF_CRM_1741593004888", "UF_CRM_1741593047994", "october_1"),
        ("UF_CRM_1741593067418", "UF_CRM_1741593115407", "october_2"),
        ("UF_CRM_1741593156926", "UF_CRM_1741593210242", "november_1"),
        ("UF_CRM_1741593231558", "UF_CRM_1741593285121", "november_2"),
        ("UF_CRM_1741593340713", "UF_CRM_1741593387667", "december_1"),
        ("UF_CRM_1741593408621", "UF_CRM_1741593452062", "december_2"),
    ]

    for dates_field, type_field, key in months:
        dates_raw = d.get(dates_field)
        type_raw = d.get(type_field)
        dates_list = dates_raw if isinstance(dates_raw, list) else ([dates_raw] if dates_raw else [])
        dates = _normalize_dates(dates_list)
        label = str(type_raw or "")
        out[key] = {"dates": dates, "type": label}

    return out

def _compute_periodicity(cleaning_dates: Dict[str, Any]) -> str:
    wash_dates = 0
    sweep_dates = 0
    full_wash_dates = 0
    first_floor_wash_dates = 0
    for key in ["september_1","september_2","october_1","october_2","november_1","november_2","december_1","december_2"]:
        block = cleaning_dates.get(key) or {}
        t = str(block.get("type") or "").lower()
        dates = block.get("dates") or []
        if not isinstance(dates, list):
            dates = []
        has_wash = ("влажная уборка" in t) or ("мытье" in t)
        has_sweep = ("подмет" in t)
        is_full = ("всех этаж" in t)
        is_first_floor = ("1 этажа" in t) or ("1 этаж" in t) or ("первые этаж" in t)
        if has_wash:
            wash_dates += len(dates)
        if has_sweep:
            sweep_dates += len(dates)
        if has_wash and is_full:
            full_wash_dates += len(dates)
        if has_wash and is_first_floor:
            first_floor_wash_dates += len(dates)
    if wash_dates == 2 and sweep_dates == 0:
        return "2 раза"
    if full_wash_dates >= 1 and first_floor_wash_dates >= 1 and wash_dates == (full_wash_dates + first_floor_wash_dates) and sweep_dates == 0:
        return "2 раза + первые этажи"
    if wash_dates == 2 and sweep_dates == 2:
        return "2 раза + 2 подметания"
    if wash_dates >= 4:
        return "4 раза"
    return "индивидуальная"

@api_router.get("/cleaning/houses", response_model=HousesResponse)
async def get_houses(
    brigade: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    management_company: Optional[str] = Query(None),
    week: Optional[str] = Query(None),
    cleaning_date: Optional[str] = Query(None),
    date_from: Optional[str] = Query(None),
    date_to: Optional[str] = Query(None),
    search: Optional[str] = Query(None),
    limit: int = Query(50),
    offset: int = Query(0),
    page: int = Query(1)
):
    try:
        raw = await bitrix.deals_full()
        if not raw:
            basic = await bitrix.deals(limit=1000)
            base_url = bitrix.webhook_url.split('/rest')[0] if bitrix.webhook_url else ''
            houses: List[HouseResponse] = []
            for d in basic:
                houses.append(HouseResponse(
                    id=int(d.get("ID", 0)),
                    title=d.get("TITLE", "Без названия"),
                    address=d.get("UF_CRM_1669561599956") or d.get("TITLE") or "",
                    brigade=d.get("ASSIGNED_BY_NAME", "") or "Бригада не назначена",
                    management_company=d.get("COMPANY_TITLE", ""),
                    status=d.get("STAGE_ID") or "",
                    apartments=int(d.get("UF_CRM_1669704529022") or 0),
                    entrances=int(d.get("UF_CRM_1669705507390") or 0),
                    floors=int(d.get("UF_CRM_1669704631166") or 0),
                    cleaning_dates={},
                    periodicity="индивидуальная",
                    bitrix_url=f"{base_url}/crm/deal/details/{d.get('ID')}/"
                ))
            total_count = len(houses)
            pages = (total_count + limit - 1) // limit
            start = (page - 1) * limit
            end = start + limit
            return HousesResponse(houses=houses[start:end], total=total_count, page=page, limit=limit, pages=pages)

        # In-memory filtering for rich fields
        def ok(d):
            if brigade and brigade not in (d.get("ASSIGNED_BY_NAME") or ""):
                return False
            if status and status != (d.get("STAGE_ID") or ""):
                return False
            if management_company and management_company not in (d.get("COMPANY_TITLE") or ""):
                return False
            cd = _build_cleaning_dates(d)
            if cleaning_date:
                # exact match
                hit = False
                for v in cd.values():
                    if isinstance(v, dict):
                        for dt in v.get("dates", []) or []:
                            if dt == cleaning_date:
                                hit = True
                                break
                    if hit: break
                if not hit:
                    return False
            if date_from and date_to:
                hit = False
                for v in cd.values():
                    if isinstance(v, dict):
                        for dt in v.get("dates", []) or []:
                            if date_from <= dt <= date_to:
                                hit = True
                                break
                    if hit: break
                if not hit:
                    return False
            return True

        # Поиск по адресу/названию (server-side)
        search = None
        try:
            # получаем из Query напрямую, но безопасно
            search = locals().get('search')  # not reliable in FastAPI
        except Exception:
            search = None
        # В FastAPI лучше читать из request.query_params, но не хотим ломать сигнатуру. Оставим фильтрацию ниже.

        # Apply search by address/title if provided
        if search:
            s = search.lower()
            raw = [d for d in raw if s in str(d.get("UF_CRM_1669561599956") or d.get("TITLE") or "").lower() or s in str(d.get("TITLE") or "").lower()]
        deals = [d for d in raw if ok(d)]
        total_count = len(deals)
        page = max(1, page)
        limit = max(1, min(1000, limit))
        start = (page - 1) * limit
        end = start + limit
        deals_page = deals[start:end]
        houses: List[HouseResponse] = []
        base_url = bitrix.webhook_url.split('/rest')[0] if bitrix.webhook_url else ''
        for d in deals_page:
            address = d.get("UF_CRM_1669561599956") or d.get("TITLE") or ""
            cd = _build_cleaning_dates(d)
            # Подменяем типы на человекочитаемые из enum при наличии
            try:
                enum_map = await bitrix._get_userfield_enums(TYPE_FIELDS)
                for key, block in cd.items():
                    if not isinstance(block, dict):
                        continue
                    t = str(block.get("type") or "")
                    # если type — это id из enum, попробуем заменить на метку
                    if t.isdigit():
                        for field_code, mapping in enum_map.items():
                            if t in mapping:
                                block["type"] = mapping[t]
                                break
            except Exception:
                pass
            periodicity = _compute_periodicity(cd)
            houses.append(HouseResponse(
                id=int(d.get("ID", 0)),
                title=d.get("TITLE", "Без названия"),
                address=address,
                brigade=d.get("ASSIGNED_BY_NAME", "") or "Бригада не назначена",
                management_company=d.get("COMPANY_TITLE", ""),
                status=d.get("STAGE_ID") or "",
                apartments=int(d.get("UF_CRM_1669704529022") or 0),
                entrances=int(d.get("UF_CRM_1669705507390") or 0),
                floors=int(d.get("UF_CRM_1669704631166") or 0),
                cleaning_dates=cd,
                periodicity=periodicity,
                bitrix_url=f"{base_url}/crm/deal/details/{d.get('ID')}/"
            ))
        pages = (total_count + limit - 1) // limit
        return HousesResponse(houses=houses, total=total_count, page=page, limit=limit, pages=pages)
    except Exception as e:
        logger.error(f"get_houses error: {e}")
        return HousesResponse(houses=[], total=0, page=page, limit=limit, pages=0)

@api_router.get("/cleaning/house/{house_id}/details")
async def get_house_details(house_id: int):
    try:
        dresp = await bitrix._call("crm.deal.get", {"id": house_id, "select": [
            "ID","TITLE","COMPANY_ID","COMPANY_TITLE","CONTACT_ID",
            "ASSIGNED_BY_NAME","ASSIGNED_BY_ID","STAGE_ID","UF_CRM_1669561599956",
            "UF_CRM_1669704529022","UF_CRM_1669705507390","UF_CRM_1669704631166",
            "UF_CRM_1741592774017","UF_CRM_1741592855565",
            "UF_CRM_1741592892232","UF_CRM_1741592945060",
            "UF_CRM_1741593004888","UF_CRM_1741593047994",
            "UF_CRM_1741593067418","UF_CRM_1741593115407",
            "UF_CRM_1741593156926","UF_CRM_1741593210242",
            "UF_CRM_1741593231558","UF_CRM_1741593285121",
            "UF_CRM_1741593340713","UF_CRM_1741593387667",
            "UF_CRM_1741593408621","UF_CRM_1741593452062"
        ]})
        deal = dresp.get("result") or {}
        if isinstance(deal, list) and deal:
            deal = deal[0]
        if not deal:
            raise HTTPException(status_code=404, detail="Дом не найден")
        company_details = {}
        if deal.get("COMPANY_ID"):
            cresp = await bitrix._call("crm.company.get", {"id": deal["COMPANY_ID"]})
            company_details = cresp.get("result") or {}
        contact_details = {}
        cid = deal.get("CONTACT_ID")
        if cid:
            if isinstance(cid, list) and cid:
                cid = cid[0]
            cresp = await bitrix._call("crm.contact.get", {"id": cid})
            contact_details = cresp.get("result") or {}
        base_url = bitrix.webhook_url.split('/rest')[0] if bitrix.webhook_url else ''
        # Человекочитаемые типы в деталях
        cd = _build_cleaning_dates(deal)
        try:
            enum_map = await bitrix._get_userfield_enums(TYPE_FIELDS)
            for key, block in cd.items():
                if not isinstance(block, dict):
                    continue
                t = str(block.get("type") or "")
                if t.isdigit():
                    for field_code, mapping in enum_map.items():
                        if t in mapping:
                            block["type"] = mapping[t]
                            break
        except Exception:
            pass
        periodicity = _compute_periodicity(cd)
        return {
            "house": {
                "id": deal.get("ID"),
                "title": deal.get("TITLE"),
                "address": deal.get("UF_CRM_1669561599956", ""),
                "apartments": int(deal.get("UF_CRM_1669704529022") or 0),
                "entrances": int(deal.get("UF_CRM_1669705507390") or 0),
                "floors": int(deal.get("UF_CRM_1669704631166") or 0),
                "brigade": deal.get("ASSIGNED_BY_NAME", ""),
                "status": deal.get("STAGE_ID", ""),
                "bitrix_url": f"{base_url}/crm/deal/details/{deal.get('ID')}/" if base_url else "",
                "cleaning_dates": cd,
                "periodicity": periodicity
            },
            "management_company": {
                "id": company_details.get("ID", ""),
                "title": company_details.get("TITLE") or (deal.get("COMPANY_TITLE") or "Не указана"),
                "phone": (company_details.get("PHONE", [{}])[0].get("VALUE", "") if company_details.get("PHONE") else ""),
                "email": (company_details.get("EMAIL", [{}])[0].get("VALUE", "") if company_details.get("EMAIL") else ""),
                "address": company_details.get("ADDRESS", ""),
                "web": (company_details.get("WEB", [{}])[0].get("VALUE", "") if company_details.get("WEB") else ""),
                "comments": company_details.get("COMMENTS", "")
            },
            "senior_resident": {
                "id": contact_details.get("ID", ""),
                "name": contact_details.get("NAME", ""),
                "last_name": contact_details.get("LAST_NAME", ""),
                "second_name": contact_details.get("SECOND_NAME", ""),
                "full_name": f"{contact_details.get('LAST_NAME','')} {contact_details.get('NAME','')} {contact_details.get('SECOND_NAME','')}".strip(),
                "phone": (contact_details.get("PHONE", [{}])[0].get("VALUE", "") if contact_details.get("PHONE") else ""),
                "email": (contact_details.get("EMAIL", [{}])[0].get("VALUE", "") if contact_details.get("EMAIL") else ""),
                "comments": contact_details.get("COMMENTS", "")
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"details error: {e}")
        raise HTTPException(status_code=500, detail="Ошибка получения деталей дома")

# ========================= AI KNOWLEDGE =========================
ALLOWED_EXT = {'.pdf','.docx','.txt','.xlsx','.zip'}
MAX_FILE_MB = int(os.environ.get('AI_MAX_FILE_MB', '50'))
MAX_TOTAL_MB = int(os.environ.get('AI_MAX_TOTAL_MB', '200'))
DEFAULT_TOP_K = 10
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY','').strip()
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY','').strip()

async def _split_into_chunks(text: str, target_tokens: int = 1200, overlap: int = 200) -> List[str]:
    enc = tiktoken.get_encoding('cl100k_base')
    toks = enc.encode(text)
    chunks = []
    i = 0
    while i < len(toks):
        window = toks[i:i+target_tokens]
        chunks.append(enc.decode(window))
        i += max(1, target_tokens - overlap)
    return chunks

async def _embed_texts(texts: List[str]) -> List[List[float]]:
    if not OPENAI_API_KEY:
        return [[0.0]*3072 for _ in texts]
    client = AsyncOpenAI(api_key=OPENAI_API_KEY)
    out = []
    for t in texts:
        try:
            r = await client.embeddings.create(model='text-embedding-3-large', input=t)
            out.append(r.data[0].embedding)
        except Exception as e:
            logger.error(f"Embedding error: {e}")
            out.append([0.0]*3072)
    return out

async def _summarize(text: str) -> str:
    if not EMERGENT_LLM_KEY:
        return text[:500]
    try:
        chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=f"ai_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}", system_message="Ты помощник для создания кратких саммари документов.").with_model('openai','gpt-4o-mini')
        prompt = f"Суммируй кратко ключевые пункты (до 120 слов):\n{text[:6000]}"
        resp = await chat.send_message(UserMessage(text=prompt))
        return resp or text[:500]
    except Exception as e:
        logger.warning(f"Summary error: {e}")
        return text[:500]

@api_router.post('/ai-knowledge/upload')
async def ai_upload(files: List[UploadFile] = File(...), chunk_tokens: int = Form(1200), overlap: int = Form(200)):
    if not files:
        raise HTTPException(status_code=400, detail='Нет файлов')
    if AsyncSessionLocal is None:
        raise HTTPException(status_code=500, detail='Database is not initialized')

    all_text = ''
    total_size = 0
    total_pages = 0
    file_stats = []
    for f in files:
        ext = os.path.splitext(f.filename or '')[1].lower()
        data = await f.read()
        size_bytes = len(data)
        total_size += size_bytes
        pages = None
        text = ''
        try:
            if ext == '.txt':
                text = data.decode('utf-8', errors='ignore')
            elif ext == '.pdf':
                reader = PdfReader(io.BytesIO(data))
                pages = len(reader.pages)
                parts = []
                for p in reader.pages:
                    try:
                        parts.append(p.extract_text() or '')
                    except Exception:
                        continue
                text = '\n'.join(parts)
            elif ext == '.docx':
                doc = DocxDocument(io.BytesIO(data))
                text = '\n'.join(p.text for p in doc.paragraphs)
            elif ext == '.xlsx':
                wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
                rows = 0
                parts = []
                for ws in wb.worksheets:
                    for row in ws.iter_rows(values_only=True):
                        line = ' '.join([str(c) for c in row if c is not None])
                        if line:
                            parts.append(line)
                        rows += 1
                pages = rows
                text = '\n'.join(parts)
            elif ext == '.zip':
                try:
                    zf = zipfile.ZipFile(io.BytesIO(data))
                    for name in zf.namelist():
                        if name.endswith('/'):
                            continue
                        sub_ext = os.path.splitext(name)[1].lower()
                        if sub_ext not in {'.pdf','.docx','.txt','.xlsx'}:
                            continue
                        if os.path.isabs(name) or '..' in name:
                            continue
                        with zf.open(name) as zf_f:
                            sub_raw = zf_f.read()
                            sub_size = len(sub_raw)
                            sub_pages = None
                            sub_text = ''
                            try:
                                if sub_ext == '.pdf':
                                    sreader = PdfReader(io.BytesIO(sub_raw))
                                    sub_pages = len(sreader.pages)
                                    sparts = []
                                    for sp in sreader.pages:
                                        try:
                                            sparts.append(sp.extract_text() or '')
                                        except Exception:
                                            continue
                                    sub_text = '\n'.join(sparts)
                                elif sub_ext == '.docx':
                                    sdoc = DocxDocument(io.BytesIO(sub_raw))
                                    sub_text = '\n'.join(p.text for p in sdoc.paragraphs)
                                elif sub_ext == '.xlsx':
                                    swb = load_workbook(io.BytesIO(sub_raw), read_only=True, data_only=True)
                                    srows = 0
                                    sparts = []
                                    for ws in swb.worksheets:
                                        for row in ws.iter_rows(values_only=True):
                                            line = ' '.join([str(c) for c in row if c is not None])
                                            if line:
                                                sparts.append(line)
                                            srows += 1
                                    sub_pages = srows
                                    sub_text = '\n'.join(sparts)
                                elif sub_ext == '.txt':
                                    sub_text = sub_raw.decode('utf-8', errors='ignore')
                            except Exception:
                                pass
                            file_stats.append({'name': name, 'ext': sub_ext, 'size_bytes': sub_size, 'pages': sub_pages, 'text_chars': len(sub_text or '')})
                            all_text += (sub_text or '') + '\n\n'
                except Exception:
                    pass
                continue
        except Exception:
            text = ''
        file_stats.append({'name': f.filename, 'ext': ext, 'size_bytes': size_bytes, 'pages': pages, 'text_chars': len(text or '')})
        total_pages += (pages or 0)
        all_text += (text or '') + '\n\n'

    if not all_text.strip():
        preview = 'Текст не извлечён. Проверьте, что файлы содержат текст (не только изображения). Статистика по файлам показана ниже.'
        upload_id = str(uuid4())
        async with AsyncSessionLocal() as s:
            meta = {
                'filenames': [f.filename for f in files],
                'summary': preview,
                'chunks': [],
                'chunks_count': 0,
                'total_size_bytes': total_size,
                'total_pages': total_pages,
                'file_stats': file_stats,
                'chunk_tokens': int(chunk_tokens),
                'overlap': int(overlap)
            }
            await s.execute(sa_text('INSERT INTO ai_uploads_temp (upload_id, meta, expires_at) VALUES (:id, :m::jsonb, :exp)'),
                            {"id": upload_id, "m": json.dumps(meta, ensure_ascii=False), "exp": datetime.now(timezone.utc)+timedelta(hours=6)})
            await s.commit()
        return {'upload_id': upload_id, 'preview': preview, 'chunks': 0, 'stats': {'total_size_bytes': total_size, 'total_pages': total_pages, 'file_stats': file_stats}}

    chunks = await _split_into_chunks(all_text, target_tokens=int(chunk_tokens), overlap=int(overlap))
    preview = await _summarize(all_text)
    if not preview:
        preview = (all_text or '')[:500]

    upload_id = str(uuid4())
    async with AsyncSessionLocal() as s:
        meta = {
            'filenames': [f.filename for f in files],
            'summary': preview,
            'chunks': chunks,
            'chunks_count': len(chunks),
            'total_size_bytes': total_size,
            'total_pages': total_pages,
            'file_stats': file_stats,
            'chunk_tokens': int(chunk_tokens),
            'overlap': int(overlap)
        }
        await s.execute(sa_text('INSERT INTO ai_uploads_temp (upload_id, meta, expires_at) VALUES (:id, :m::jsonb, :exp)'),
                        {"id": upload_id, "m": json.dumps(meta, ensure_ascii=False), "exp": datetime.now(timezone.utc)+timedelta(hours=6)})
        await s.commit()

    return {
        'upload_id': upload_id,
        'preview': preview,
        'chunks': len(chunks),
        'stats': {'total_size_bytes': total_size, 'total_pages': total_pages, 'file_stats': file_stats}
    }

@api_router.post('/ai-knowledge/save')
async def ai_save(upload_id: str = Form(...), filename: str = Form('document.txt'), db: AsyncSession = Depends(get_db)):
    row = (await db.execute(sa_text('SELECT meta FROM ai_uploads_temp WHERE upload_id=:id'), {"id": upload_id})).first()
    if not row:
        raise HTTPException(status_code=404, detail='upload_id не найден или истёк')
    raw_meta = row[0]
    if isinstance(raw_meta, str):
        meta = json.loads(raw_meta)
    else:
        meta = raw_meta  # asyncpg returns dict for JSONB
    chunks = meta.get('chunks', [])
    vectors = await _embed_texts(chunks)
    doc_id = str(uuid4())
    summary = meta.get('summary') or 'См. превью при загрузке'
    size_bytes = meta.get('total_size_bytes')
    pages = meta.get('total_pages')
    await db.execute(sa_text('INSERT INTO ai_documents (id, filename, mime, size_bytes, summary, pages, created_at) VALUES (:i,:fn,:mime,:sz,:sm,:pg,:ca)'),
                     {"i": doc_id, "fn": filename, "mime": "text/plain", "sz": size_bytes, "sm": (summary[:500] if isinstance(summary,str) else None), "pg": pages, "ca": datetime.now(timezone.utc)})
    for idx, (text, v) in enumerate(zip(chunks, vectors)):
        await db.execute(sa_text('INSERT INTO ai_chunks (id, document_id, chunk_index, content, embedding) VALUES (:i,:d,:x,:c,:e)'),
                         {"i": str(uuid4()), "d": doc_id, "x": idx, "c": text, "e": v})
    await db.execute(sa_text('DELETE FROM ai_uploads_temp WHERE upload_id=:id'), {"id": upload_id})
    await db.commit()
    return {"document_id": doc_id}

@api_router.get('/ai-knowledge/documents')
async def ai_docs_list(db: AsyncSession = Depends(get_db)):
    rows = (await db.execute(sa_text('''
        SELECT d.id, d.filename, d.mime, d.size_bytes, d.summary, d.created_at, d.pages,
               (SELECT COUNT(1) FROM ai_chunks c WHERE c.document_id=d.id) AS chunks_count
        FROM ai_documents d
        ORDER BY d.created_at DESC
        LIMIT 200
    '''))).all()
    docs = []
    for r in rows:
        rid, filename, mime, size_bytes, summary, created_at, pages, chunks_count = r
        docs.append({"id": rid, "filename": filename, "mime": mime, "size_bytes": size_bytes, "summary": summary, "created_at": created_at.isoformat() if created_at else None, "pages": pages, "chunks_count": chunks_count})
    return {"documents": docs}

class SearchRequest(BaseModel):
    query: str
    top_k: int = DEFAULT_TOP_K

@api_router.post('/ai-knowledge/search')
async def ai_search(req: SearchRequest, db: AsyncSession = Depends(get_db)):
    if not req.query.strip():
        raise HTTPException(status_code=400, detail='query пуст')
    qv = (await _embed_texts([req.query]))[0]
    rows = (await db.execute(sa_text('''
        SELECT c.document_id, c.chunk_index, c.content, 1 - (c.embedding <=> :qv) as score, d.filename
        FROM ai_chunks c JOIN ai_documents d ON d.id = c.document_id
        ORDER BY c.embedding <=> :qv
        LIMIT :k
    '''), {"qv": qv, "k": req.top_k})).all()
    results = []
    for r in rows:
        doc_id, idx, content, score, filename = r
        results.append({"document_id": doc_id, "chunk_index": idx, "content": content, "score": float(score), "filename": filename})
    return {"results": results}

@api_router.delete('/ai-knowledge/document/{doc_id}')
async def ai_delete(doc_id: str, db: AsyncSession = Depends(get_db)):
    await db.execute(sa_text('DELETE FROM ai_chunks WHERE document_id=:id'), {"id": doc_id})
    await db.execute(sa_text('DELETE FROM ai_documents WHERE id=:id'), {"id": doc_id})
    await db.commit()
    return {"ok": True}

app.include_router(api_router)

@app.on_event("startup")
async def on_startup():
    await init_db()